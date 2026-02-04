"""Generate a printable Word document for the Lab 2 Answer Sheet."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_answer_lines(doc, num_lines=4):
    """Add blank ruled lines for student answers."""
    for _ in range(num_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(24)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

def add_bold_text(paragraph, text):
    """Add bold text to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_normal_text(paragraph, text):
    """Add normal text to a paragraph."""
    run = paragraph.add_run(text)
    return run

def add_code_text(paragraph, text):
    """Add monospace code-styled text to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return run

def set_narrow_margins(section):
    """Set narrow margins for more printable space."""
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def build_document():
    doc = Document()

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # -- Heading styles --
    for level in [1, 2, 3]:
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(0)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(2)

    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(8)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(4)

    doc.styles['Heading 3'].font.size = Pt(11)
    doc.styles['Heading 3'].paragraph_format.space_before = Pt(6)
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(2)

    # -- Margins --
    section = doc.sections[0]
    set_narrow_margins(section)

    # ===== PAGE 1: Title + Module 0 =====

    # Title block
    h = doc.add_heading('Lab 2: Gradient Descent', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.add_heading('Answer Sheet', level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bold_text(p, 'DATA 1010 -- AI in Action')

    # Name / Group Code line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    add_bold_text(p, 'Name(s): ')
    add_normal_text(p, '______________________________________     ')
    add_bold_text(p, 'Group Code: ')
    add_normal_text(p, '___________')

    add_horizontal_rule(doc)

    # Overview box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    add_bold_text(p, 'Key Concept: ')
    add_normal_text(p, 'Gradient descent uses the rule ')
    add_code_text(p, 'new = old - learning_rate \u00d7 gradient')
    add_normal_text(p, ' to automatically navigate toward minima.')

    p = doc.add_paragraph()
    add_bold_text(p, 'Lab Structure: ')
    add_normal_text(p, '5 modules (0\u20134) using the same group code throughout.')

    # ----- Module 0 -----
    add_horizontal_rule(doc)
    doc.add_heading('Module 0: Setup & The Update Rule (~5 min)', level=2)

    p = doc.add_paragraph()
    add_bold_text(p, 'Concepts:')
    bullets = [
        ('Universal update rule: ', 'new = old - learning_rate \u00d7 gradient', True),
        ('Gradient', ' = local slope = direction of steepest ascent', False),
        ('Learning rate', ' = step size multiplier', False),
        ('Move downhill', ' = negate gradient', False),
    ]
    for bold_part, rest, is_code in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_bold_text(bp, bold_part)
        if is_code:
            add_code_text(bp, rest)
        else:
            add_normal_text(bp, rest)

    # Q1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q1. ')
    add_normal_text(p, 'If the gradient (slope) at a point is ')
    add_bold_text(p, 'positive')
    add_normal_text(p, ', which direction does gradient descent move? Why?')
    add_answer_lines(doc, 3)

    # Q2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q2. ')
    add_normal_text(p, 'What happens to the step size if:')
    for item in ['(a) The learning rate is very large (e.g., 10.0)?',
                 '(b) The gradient magnitude is very large?',
                 '(c) Both learning rate and gradient are small?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_normal_text(bp, item)
    add_answer_lines(doc, 4)

    # ===== PAGE 2: Module 1 =====
    add_page_break(doc)

    doc.add_heading('Module 1: GD on Hidden Parabola (~15 min)', level=2)

    p = doc.add_paragraph()
    add_bold_text(p, 'What you do: ')
    add_normal_text(p, 'Apply GD to a hidden 1D function. Watch it automatically find the minimum.')

    # Q3
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q3 \u2013 PREDICTION: ')
    add_normal_text(p, 'Starting from x = 0.0, predict what will happen with five learning rates:')
    for item in ['LR = 0.01: Will this converge quickly or slowly?',
                 'LR = 0.05: Faster than 0.01?',
                 'LR = 0.4: Any risks?',
                 'LR = 1.0: Fastest, or problems?',
                 'LR = 3.0: What do you expect?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        # Bold the LR part
        parts = item.split(':')
        add_bold_text(bp, parts[0] + ':')
        add_normal_text(bp, parts[1])

    doc.add_heading('Prediction:', level=3)
    add_answer_lines(doc, 4)

    doc.add_heading('Result after running:', level=3)
    add_answer_lines(doc, 4)

    # Q4
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q4. ')
    add_normal_text(p, 'Based on the visualizations:')
    for item in ['How does step size relate to (a) gradient magnitude and (b) learning rate?',
                 'Why do steps get smaller near the minimum?',
                 'Describe what happens with LR = 1.0 and LR = 3.0.']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_normal_text(bp, item)
    add_answer_lines(doc, 5)

    # ===== PAGE 3: Module 2 =====
    add_page_break(doc)

    doc.add_heading('Module 2: GD on Parameter Space \u2013 Line Fitting (~20 min)', level=2)

    p = doc.add_paragraph()
    add_bold_text(p, 'What you do: ')
    add_normal_text(p, 'Apply GD to optimize two parameters (m, b) simultaneously. Watch GD navigate the MSE landscape.')

    # Q5
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q5 \u2013 PREDICTION: ')
    add_normal_text(p, 'Starting from (0, 0), predict: will the GD path be straight or curved? Why?')

    doc.add_heading('Prediction:', level=3)
    add_answer_lines(doc, 3)

    doc.add_heading('Result:', level=3)
    add_answer_lines(doc, 3)

    # Q6
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q6. ')
    add_normal_text(p, 'Describe the GD path and the learning rate comparison:')
    for item in ['Is the path straight or curved? Why?',
                 'What happens to step size near the minimum?',
                 'Which learning rate (0.01, 0.1, 0.5) converged fastest, and what went wrong with the others?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_normal_text(bp, item)
    add_answer_lines(doc, 5)

    # ===== PAGE 4: Module 3 =====
    add_page_break(doc)

    doc.add_heading('Module 3: Learning Rate Exploration (~20 min)', level=2)

    p = doc.add_paragraph()
    add_bold_text(p, 'What you do: ')
    add_normal_text(p, 'Deep dive into learning rate effects using a simple function. Run GD with LR = {0.001, 0.1, 0.8, 3.0}.')

    # Q7
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q7 \u2013 PREDICTION: ')
    add_normal_text(p, 'Starting from x = 10.0, predict for each learning rate:')
    for item in ['LR = 0.001 (very small): Will it converge in 100 steps?',
                 'LR = 0.1 (moderate): Fast convergence?',
                 'LR = 0.8 (large): Converge, oscillate, or diverge?',
                 'LR = 3.0 (very large): What do you expect?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        parts = item.split(':')
        add_bold_text(bp, parts[0] + ':')
        add_normal_text(bp, ':'.join(parts[1:]))

    doc.add_heading('Prediction:', level=3)
    add_answer_lines(doc, 4)

    doc.add_heading('Result after running:', level=3)
    add_answer_lines(doc, 4)

    # Q8
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q8. ')
    add_normal_text(p, 'Describe the behavior for each LR category:')
    for item in ['Too small (0.001): What happens? Why is this wasteful?',
                 'Just right (0.1): What makes this work well?',
                 'Too large (0.8): What problems occur?',
                 'Way too large (3.0): What does divergence look like?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        parts = item.split(':')
        add_bold_text(bp, parts[0] + ':')
        add_normal_text(bp, ':'.join(parts[1:]))
    add_answer_lines(doc, 5)

    # Q9
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q9. ')
    add_normal_text(p, 'How would you choose a learning rate for a new problem? What signs tell you it\'s too large or too small?')
    add_answer_lines(doc, 4)

    # ===== PAGE 5: Module 4 + Key Takeaways =====
    add_page_break(doc)

    doc.add_heading('Module 4: Mountain Landscape \u2013 GD Limitations (~15 min)', level=2)

    p = doc.add_paragraph()
    add_bold_text(p, 'What you do: ')
    add_normal_text(p, 'Run gradient ascent (uphill climbing) from multiple starting points on a landscape with several peaks.')

    # Q10
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q10 \u2013 PREDICTION: ')
    add_normal_text(p, 'Before running gradient ascent:')
    for item in ['Starting at (1, 1): Will GD find the global maximum? Why or why not?',
                 'Will different starting points reach different peaks?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_normal_text(bp, item)

    doc.add_heading('Prediction:', level=3)
    add_answer_lines(doc, 3)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    add_bold_text(p, 'Result after running: ')
    add_normal_text(p, 'How many different peaks did you reach from different starting points?')
    add_answer_lines(doc, 2)

    # Q11
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Q11. ')
    add_normal_text(p, 'Based on your experiments:')
    for item in ['Did gradient ascent find the global maximum from every starting point?',
                 'Why can\'t GD "see" distant peaks?',
                 'What strategies might help overcome this limitation?']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(1)
        add_normal_text(bp, item)
    add_answer_lines(doc, 5)

    # Key Takeaways
    add_horizontal_rule(doc)
    doc.add_heading('Key Takeaways', level=2)

    takeaways = [
        ('Universal update rule: ', 'new = old - learning_rate \u00d7 gradient', ' works for any optimization'),
        ('GD automates search: ', 'Replaces manual exploration with systematic gradient-following', ''),
        ('Learning rate is critical: ', 'Too small = slow, too large = unstable, "just right" = optimal', ''),
        ('Local optima problem: ', 'GD gets stuck at the first peak/valley it reaches', ''),
        ('Starting point matters: ', 'Different initializations lead to different solutions', ''),
    ]
    for bold_part, text, suffix in takeaways:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.space_before = Pt(2)
        add_bold_text(bp, bold_part)
        add_normal_text(bp, text + suffix)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_bold_text(p, 'Connection to ML: ')
    add_normal_text(p, 'Everything you learned applies to training neural networks with millions of parameters navigating complex loss landscapes!')

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Lab_2_Answer_Sheet.docx')
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path

if __name__ == '__main__':
    build_document()
